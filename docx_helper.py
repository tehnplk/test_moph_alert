#!/usr/bin/env python3
"""
ฟังก์ชันช่วยสร้างเอกสาร DOCX ภาษาไทยตามรูปแบบงานสารบรรณ

ใช้ร่วมกันระหว่างสคริปต์สร้างเอกสารนวัตกรรมทุกฉบับ
แก้ฟอนต์หรือรูปแบบที่นี่ที่เดียว มีผลกับทุกเอกสาร
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "TH SarabunPSK"
BODY_SIZE = 16
TABLE_SIZE = 14

NAVY = RGBColor(0x1F, 0x3B, 0x73)

BLANK = "…………………………."      # ช่องเว้นว่างยาว
BLANK_S = "…………"              # ช่องเว้นว่างสั้น (ในตาราง)

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ──────────────────────────────────────────────────────────────
# ขอบเขตการแจ้งข่าวการระบาด จำแนกตามโรค
#
# ใช้ร่วมกันทั้งฉบับ รพ.สต. และฉบับ สสจ. เพราะเป็นแนวทางเดียวกัน
# ประเด็นสำคัญคือแต่ละโรคแพร่คนละแบบ ขอบเขตกลุ่มเป้าหมายจึงต่างกัน
# ──────────────────────────────────────────────────────────────
DISEASE_ROWS = [
    ["โรคติดต่อ", "การกำหนดกลุ่มเป้าหมาย", "สาระสำคัญของข้อความ"],
    ["ไข้เลือดออก",
     "ครัวเรือนในรัศมี ๑๐๐ เมตร\nรอบบ้านผู้ป่วย ตามระยะบินของยุงลาย",
     "มาตรการ ๓ เก็บ ๓ ป้องกัน\nและอาการที่ต้องรีบพบแพทย์"],
    ["เลปโตสไปโรซิส\n(โรคฉี่หนู)",
     "ครัวเรือนในพื้นที่น้ำท่วมขัง\nและผู้ประกอบอาชีพเกษตรกรรม",
     "หลีกเลี่ยงการเดินลุยน้ำ สวมรองเท้าบูท\nและอาการไข้ร่วมกับปวดกล้ามเนื้อน่อง"],
    ["โรคติดเชื้อไวรัสโคโรนา ๒๐๑๙",
     "ผู้สัมผัสใกล้ชิด และประชาชน\nในพื้นที่ที่พบการระบาดเป็นกลุ่มก้อน",
     "การป้องกันตนเอง การตรวจหาเชื้อ\nและการดูแลกลุ่มเสี่ยง ๖๐๘"],
    ["อาหารเป็นพิษ",
     "ผู้ร่วมรับประทานอาหารในงาน\nหรือสถานที่เดียวกันกับผู้ป่วย",
     "อาการที่ต้องเฝ้าระวัง การดื่มสารละลาย\nเกลือแร่ และการงดอาหารที่สงสัย"],
    ["อุจจาระร่วงเฉียบพลัน",
     "ครัวเรือนที่ใช้แหล่งน้ำร่วมกัน\nศูนย์พัฒนาเด็กเล็ก และโรงเรียน",
     "กินร้อน ช้อนกลาง ล้างมือ\nและการดูแลผู้ป่วยเบื้องต้น"],
    ["โรคติดต่ออื่น",
     "ตามแนวทางการสอบสวนโรค\nและควบคุมโรคของแต่ละโรค",
     "ตามคำแนะนำของกรมควบคุมโรค"],
]

DISEASE_WIDTHS = [4.0, 6.2, 6.3]


def thai(run, size=BODY_SIZE, bold=False, color=None):
    """ตั้งฟอนต์ไทยให้ครบทั้ง ascii / eastAsia / complex-script

    ถ้าไม่ตั้ง w:cs และ w:szCs ด้วย Word จะ render ภาษาไทย
    ด้วยฟอนต์อื่นและขนาดเพี้ยนจากที่กำหนด
    """
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT)
    rpr.append(rpr.makeelement(qn("w:szCs"), {qn("w:val"): str(int(size * 2))}))
    if bold:
        rpr.append(rpr.makeelement(qn("w:bCs"), {}))
    return run


def new_document():
    """สร้างเอกสารเปล่าที่ตั้งขอบกระดาษและฟอนต์เริ่มต้นไว้แล้ว"""
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn("w:cs"), FONT)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return doc


def para(doc, text="", size=BODY_SIZE, bold=False, align=None, space_after=6,
         indent=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    thai(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, size=18):
    return para(doc, text, size=size, bold=True, space_after=5, color=NAVY)


def bullet(doc, text, indent=0.8):
    return para(doc, "•  " + text, indent=indent, space_after=3)


def steps(doc, items, indent=1.0):
    """รายการขั้นตอน — items เป็น list ของ (ป้ายกำกับ, ข้อความ)"""
    for label, text in items:
        para(doc, label + "   " + text, indent=indent, space_after=4)


def table(doc, rows, widths=None, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if c > 0:
                p.alignment = CENTER
            thai(p.add_run(str(val)), size=TABLE_SIZE, bold=(header and r == 0))
    if widths:
        for row in t.rows:
            for c, w in enumerate(widths):
                row.cells[c].width = Cm(w)
    return t


def cover(doc, kicker, title, subtitle=None, tagline=None, org=None):
    """หน้าปกเอกสาร"""
    para(doc, kicker, size=18, bold=True, align=CENTER, space_after=2)
    para(doc, title, size=30, bold=True, align=CENTER, space_after=4, color=NAVY)
    if subtitle:
        para(doc, subtitle, size=18, bold=True, align=CENTER, space_after=4)
    if tagline:
        para(doc, tagline, size=18, align=CENTER, space_after=2)
    if org:
        para(doc, org, size=18, bold=True, align=CENTER, space_after=16)
